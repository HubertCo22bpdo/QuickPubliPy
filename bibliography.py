from docx import Document
import json
from fetchmeta import fetchmeta
from ast import literal_eval
# import bib2json




class Bibliography:
    def __init__(self, identifier):
        self.word_file = Document()
        if type(identifier) == dict:
            self.bib = identifier
        if type(identifier) == str:
            with open(f'{identifier}', 'r') as f:
                self.bib = json.load(f)


    def insert(self, data, index=999999):
        if index > len(self.bib):

            self.bib[len(self.bib) + 1] = data
        else:
            for key in sorted(self.bib.keys(), reverse=True, key=int):
                key = int(key)
                if key >= index:
                    try:
                        self.bib[str(key + 1)] = self.bib.pop(str(key))
                    except:
                        try:
                            self.bib[key + 1] = self.bib.pop(key)
                        except e as e:
                            print(e, '\nActual indexing error')

            self.bib[index] = data
        self.bib = dict(sorted(self.bib.items(), key=lambda dictbib: int(dictbib[0])))

    def insert_from_doi(self, doi, index=999999):
        data = fetchmeta(f'{doi}')
        # with open(fr'.\_temp.bib', 'w') as f:
        #     f.write(fetchmeta(f'{doi}', 'bibtex'))
        # bib2json.convert_file(".\_temp.bib", ".\_temp.json")
        # with open(fr'.\_temp.json', 'r') as f:
        #     data = json.load(f)
        authors_str = ''
        for author in data['author']:
            authors_str += author['given'] + ' ' + author['family'] + ', '

        if 'page' in data.keys():
            if '-' in data['page']:
                pages = data['page'].split('-')
            else:
                pages = data['page']
        else:
            pages = 'Not found'
        
        formated_data = {
            'Authors': authors_str[:-2],
            'DOI': doi,
            'Year': data['published']['date-parts'][0][0],
            'Pages': pages,
            'Journal': data['container-title'],
            'Title': data['title'],
            'Volume': data['volume'] if 'issue'in data.keys() else 'Early view',
            'Issue': data['issue'] if 'issue'in data.keys() else None
        }
        self.insert(formated_data, index=index)
        
    def remove(self, index):
        try:
            del self.bib[str(index)]
        except:
            del self.bib[int(index)]
        for key in sorted(self.bib.keys(), key=int):
            key = int(key)
            if key >= index:
                try:
                    self.bib[str(key - 1)] = self.bib.pop(str(key))
                except:
                    try:
                        self.bib[key - 1] = self.bib.pop(key)
                    except e as e:
                        print(e, '\nActual indexing error')

        self.bib = dict(sorted(self.bib.items(), key=lambda dictbib: int(dictbib[0])))

    def export_bibliography(self, style, filename='defult_bib.docx'):
        for nr, info_dict in self.bib.items():
            if style == 'Chicago':
                """
                A. Surname1, B. Surname2, C. Surname3, et al., (more than 6) "Title" Journal(italic) year(bold), volume(italic), issue, pages, DOI
                """
                citation = self.word_file.add_paragraph(f'{nr}.\t')
                names_str = ''
                full_names = [name.strip() for name in info_dict["Authors"].split(',') if name.strip()]

                # Known surname particles
                particles = {'de', 'del', 'de la', 'van', 'von', 'der', 'den', 'la', 'le', 'du', 'di', 'da', 'ter'}

                def format_author(name: str) -> str:
                    parts = name.strip().split()
                    if not parts:
                        return ''
                    
                    # --- Determine last name ---
                    last_name_parts = []
                    i = len(parts) - 1
                    last_name_parts.insert(0, parts[i])
                    i -= 1
                    while i >= 0:
                        particle_candidate = parts[i].lower()
                        two_word = f'{parts[i].lower()} {parts[i+1].lower()}'
                        if i < len(parts)-1 and two_word in particles:
                            last_name_parts.insert(0, parts[i])
                            i -= 1
                        elif particle_candidate in particles:
                            last_name_parts.insert(0, parts[i])
                            i -= 1
                        else:
                            break
                    last_name = ' '.join(last_name_parts)

                    # --- Handle initials ---
                    initials_parts = parts[:i+1]
                    initials = []
                    for token in initials_parts:
                        if '-' in token:
                            subparts = token.split('-')
                            if len(subparts) == 2:
                                first, second = subparts
                                if second[0].islower():
                                    # Treat whole as one name: Shin-ichi → S.
                                    initials.append(f'{first[0]}.')
                                else:
                                    # Jun-Hao → J.-H.
                                    initials.append(f'{first[0]}.-{second[0]}.')
                            else:
                                initials.append(f'{subparts[0][0]}.')  # fallback
                        elif token.endswith('.'):
                            # Already an initial
                            initials.append(token)
                        else:
                            initials.append(f'{token[0]}.')
                    
                    return f'{" ".join(initials)} {last_name}'

                # --- Format list of names ---
                if len(full_names) > 6:
                    for author in full_names[:3]:
                        names_str += f'{format_author(author)}, '
                    names_str += 'et al., '
                else:
                    for author in full_names:
                        names_str += f'{format_author(author)}, '

                citation.add_run(names_str)

                #Title
                citation.add_run('"')
                title = info_dict["Title"]
                format_stack = []
                i = 0
                while i < len(title):
                    if title.startswith(r'\textit{', i):
                        format_stack.append('italic')
                        i += len(r'\textit{')
                    elif title.startswith(r'<i>', i):
                        format_stack.append('italic')
                        i += len(r'<i>')
                    elif title.startswith(r'\textbf{', i):
                        format_stack.append('bold')
                        i += len(r'\textbf{')
                    elif title.startswith('^{', i):
                        format_stack.append('superscript')
                        i += 2
                    elif title.startswith('<sup>', i):
                        format_stack.append('superscript')
                        i += len(r'<sup>')
                    elif title.startswith('</sup>', i):
                        if format_stack:
                            format_stack.pop()
                        i += len(r'</sup>')
                    elif title.startswith('_{', i):
                        format_stack.append('subscript')
                        i += 2
                    elif title.startswith('<sub>', i):
                        format_stack.append('subscript')
                        i += len(r'<sub>')
                    elif title.startswith('</sub>', i):
                        if format_stack:
                            format_stack.pop()
                        i += len(r'</sub>')
                    elif title.startswith('</i>', i):
                        if format_stack:
                            format_stack.pop()
                        i += len(r'</i>')
                    elif title.startswith('}\\', i):
                        if format_stack:
                            format_stack.pop()
                        i += 2
                    else:
                        run = citation.add_run(title[i])
                        run.font.italic = 'italic' in format_stack
                        run.font.bold = 'bold' in format_stack
                        run.font.superscript = 'superscript' in format_stack
                        run.font.subscript = 'subscript' in format_stack
                        i += 1
                citation.add_run('" ')

                #Journal name
                citation.add_run(f'{info_dict["Journal"]} ').italic = True

                #Year
                citation.add_run(f'{info_dict["Year"]}').bold = True
                citation.add_run(f', ')

                #Volume
                citation.add_run(f'{info_dict["Volume"]}').italic = True
                citation.add_run(f', ')

                #Issue
                if info_dict["Issue"] is not None:
                    citation.add_run(f'{info_dict["Issue"]}, ')
                
                #Pages
                if "['" in info_dict["Pages"]:
                    to_eval = literal_eval(info_dict["Pages"])
                    if (type(to_eval[0]) == int) and (type(to_eval[1]) == int):
                        info_dict["Pages"] = literal_eval(info_dict["Pages"])
                if type(info_dict["Pages"]) == tuple or type(info_dict["Pages"]) == list:
                    citation.add_run(f'{info_dict["Pages"][0]}–{info_dict["Pages"][1]},')
                else:
                    citation.add_run(f'{info_dict["Pages"]},')

                #DOI
                self.word_file.add_paragraph(f'DOI: {info_dict["DOI"]}')

                self.word_file.save(filename)
            elif style == "RSC":
                """
                A. Surname1, B. Surname2 and C. Surname3, Journal(italic) year, volume(bold), pages
                """
                citation = self.word_file.add_paragraph(f'{nr}.\t')
                names_str = ''
                full_names = [name.strip() for name in info_dict["Authors"].split(',') if name.strip()]

                # Known surname particles
                particles = {'de', 'del', 'de la', 'van', 'von', 'der', 'den', 'la', 'le', 'du', 'di', 'da', 'ter'}

                def format_author(name: str) -> str:
                    parts = name.strip().split()
                    if not parts:
                        return ''
                    
                    # --- Determine last name ---
                    last_name_parts = []
                    i = len(parts) - 1
                    last_name_parts.insert(0, parts[i])
                    i -= 1
                    while i >= 0:
                        particle_candidate = parts[i].lower()
                        two_word = f'{parts[i].lower()} {parts[i+1].lower()}'
                        if i < len(parts)-1 and two_word in particles:
                            last_name_parts.insert(0, parts[i])
                            i -= 1
                        elif particle_candidate in particles:
                            last_name_parts.insert(0, parts[i])
                            i -= 1
                        else:
                            break
                    last_name = ' '.join(last_name_parts)

                    # --- Handle initials ---
                    initials_parts = parts[:i+1]
                    initials = []
                    for token in initials_parts:
                        if '-' in token:
                            subparts = token.split('-')
                            if len(subparts) == 2:
                                first, second = subparts
                                if second[0].islower():
                                    # Treat whole as one name: Shin-ichi → S.
                                    initials.append(f'{first[0]}.')
                                else:
                                    # Jun-Hao → J.-H.
                                    initials.append(f'{first[0]}.-{second[0]}.')
                            else:
                                initials.append(f'{subparts[0][0]}.')  # fallback
                        elif token.endswith('.'):
                            # Already an initial
                            initials.append(token)
                        else:
                            initials.append(f'{token[0]}.')
                    
                    return f'{" ".join(initials)} {last_name}'

                for author_i in range(len(full_names)):
                    if (author_i == (len(full_names) - 1)):
                        names_str += f'and {format_author(full_names[author_i])}, '
                    else:
                        names_str += f'{format_author(full_names[author_i])}, '

                citation.add_run(names_str)


                #Journal name
                citation.add_run(f'{info_dict["Journal"]}, ').italic = True

                #Year
                citation.add_run(f'{info_dict["Year"]}') = True
                citation.add_run(f', ')

                #Volume
                citation.add_run(f'{info_dict["Volume"]}').bold = True
                citation.add_run(f', ')
                
                #Pages
                if "['" in info_dict["Pages"]:
                    to_eval = literal_eval(info_dict["Pages"])
                    if (type(to_eval[0]) == int) and (type(to_eval[1]) == int):
                        info_dict["Pages"] = literal_eval(info_dict["Pages"])
                if type(info_dict["Pages"]) == tuple or type(info_dict["Pages"]) == list:
                    citation.add_run(f'{info_dict["Pages"][0]}–{info_dict["Pages"][1]}')
                else:
                    citation.add_run(f'{info_dict["Pages"]}')


    def save_json(self, filename):
        self.bib = dict(sorted(self.bib.items(), key=lambda dictbib: int(dictbib[0])))
        with open(f'{filename}', 'w') as f:
            json.dump(self.bib, f, indent=2)
